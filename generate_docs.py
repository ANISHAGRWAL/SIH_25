import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = create_element('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def main():
    doc = docx.Document()

    # Define base styles or customize default normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ----------------------------------------------------
    # COVER PAGE / MAIN PLATFORM TITLE
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("CAMPUS CARE")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121) # Sleek blue
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)

    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Mental Health & Wellness Platform for Higher Education\nSystem Architecture, Detailed Workflows, & Entity Relationship Database Documentation")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle_p.paragraph_format.space_after = Pt(48)

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 1: ARCHITECTURE & WORKFLOW
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Technical Workflow & Platform Architecture")
    h1_run.font.color.rgb = RGBColor(31, 78, 121)
    h1_run.font.bold = True

    p_intro = doc.add_paragraph("Campus Care is a comprehensive, privacy-first digital mental health platform designed for students in higher education. It connects self-assessments, real-time facial/vocal emotion tracking, AI journaling analysis, anonymous peer support forums, therapist-grade AI chatbot assistance, and one-tap counselor bookings into a unified, responsive portal.")

    # Stack breakdown
    h2_stack = doc.add_heading(level=2)
    h2_stack.add_run("1.1 Technology Stack Matrix").font.color.rgb = RGBColor(46, 117, 182)
    
    bullets = [
        ("Frontend Client", "Next.js 14 (App Router) with TypeScript, styled with Tailwind CSS + ShadCN UI. Capacitor.js is integrated to compile the web application into native Android/iOS app shells."),
        ("Application Server", "Node.js & Express (TypeScript) handling user authentication (JWT), Socket.io real-time operations, and third-party integrations (SendGrid, Cloudinary, AssemblyAI)."),
        ("Database Layer", "PostgreSQL database hosted on Neon DB, using Drizzle ORM for schema generation and queries."),
        ("AI/ML Services", "FastAPI Python services: one for custom Keras CNN speech emotion classification (Port 8001) and one for LangGraph AI Chat Agent (Port 8000).")
    ]
    for role, desc in bullets:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(role + ": ")
        r_bold.bold = True
        p.add_run(desc)

    # Core workflows
    h2_work = doc.add_heading(level=2)
    h2_work.add_run("1.2 Core System Workflows (Step-by-Step)").font.color.rgb = RGBColor(46, 117, 182)

    workflows = [
        ("Authentication & Onboarding", 
         "Students sign up by choosing their college organization (validating email domains). Credentials, personal bios, and emergency contacts are filled, and an ID proof is uploaded via Multer to Cloudinary. Security is enforced with email-based OTP verification and JWT-based session authorization."),
        
        ("Facial Emotion Detection (Client-Side)",
         "Using client-side face-api.js (Tiny Face Detector and Face Expression Net), the browser captures camera frames via getUserMedia. It runs local inference to determine the user's primary expression (happy, sad, angry, surprised, disgusted, fearful, neutral) and calls the backend POST API to submit results. Customized activities and blog redirects are then loaded on-screen."),
        
        ("Vocal Mood Detection (Server-Side)",
         "Students record voice samples via browser MediaRecorder. The webm binary is posted to the FastAPI Speech service. The service downsamples audio to 22.050 Hz, trims silence, normalizes volume, extracts 40 MFCC features across 200 time-frames, and feeds the (1, 40, 200, 1) tensor to a custom Keras CNN model (my_model.h5) returning the emotion classification."),
         
        ("AI Journaling (Mind Log)",
         "Students write their personal diary logs. The backend forwards the text to Google's gemini-2.5-flash model, which extracts numeric distress scores (0 to 10) on five dimensions: Mood Disturbance, Sleep Disruption, Appetite Issues, Academic Disengagement, and Social Withdrawal. Results are stored in the database, with a rule-based parser falling back to a regex heuristic if the model fails."),
         
        ("Psych-Tests Surveys",
         "Enforces clinical screening tests (PHQ-9 for depression, GAD-7 for anxiety, PSS-10 for stress) with cooldown guards (7 days for PHQ/GAD, 30 days for PSS) to track scores chronologically and prevent survey fatigue."),
         
        ("Safety Routing & Emergency SOS",
         "All AI chatbot queries are audited by a dual classifier. If self-harm/suicidal intent is found (Crisis), the system triggers an SOS alert: it queries database user and organization admins, sends an email notification via SendGrid, and redirects the student to helplines (+91 9152987821) and the campus admin's contact. If the query requires clinical support (Therapy), it routes to a local Ollama server running MedGemma:4b. Otherwise, it triggers the standard friendly chatbot (Groq Llama-3.3 Tough Coach / Gemini Flash Comfort Bot)."),
         
        ("Anonymous Peer Support Chat",
         "A real-time socket connection connects students to campus volunteers. Students request chat sessions; volunteers receive notifications and accept requests. A room is established via Socket.io. When either user leaves, a database transaction is committed to erase all messages and session logs, ensuring complete, audit-free anonymity.")
    ]

    for title, desc in workflows:
        p_wf = doc.add_paragraph()
        r_wf_bold = p_wf.add_run(f"• {title}: ")
        r_wf_bold.bold = True
        r_wf_bold.font.color.rgb = RGBColor(31, 78, 121)
        p_wf.add_run(desc)
        p_wf.paragraph_format.space_before = Pt(4)
        p_wf.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 2: ER DIAGRAM & DATA DICTIONARY
    # ----------------------------------------------------
    h1_er = doc.add_heading(level=1)
    h1_er_run = h1_er.add_run("2. Database Entity Relationship (ER) Schema")
    h1_er_run.font.color.rgb = RGBColor(31, 78, 121)
    h1_er_run.font.bold = True

    p_er_intro = doc.add_paragraph("The database layer is managed through serverless PostgreSQL (Neon DB) using Drizzle ORM. Relations are strictly mapped using foreign key constraints with on-delete cascading rules to maintain database integrity.")

    h2_schema = doc.add_heading(level=2)
    h2_schema.add_run("2.1 Entity Attributes & Field Specifications").font.color.rgb = RGBColor(46, 117, 182)

    # We will represent the tables as clear text headings with key details.
    tables_spec = [
        ("1. organizations", "Stores campus institutional nodes.", [
            ("id", "UUID", "Primary Key (Auto-generated UUID)"),
            ("name", "VARCHAR(255)", "Institution name"),
            ("createdAt / updatedAt", "TIMESTAMP", "Auditing timestamps")
        ]),
        ("2. users", "Represents students, admins, and volunteer actors.", [
            ("id", "UUID", "Primary Key"),
            ("email", "VARCHAR(255)", "Unique login identifier"),
            ("password", "VARCHAR(255)", "Secure hashed password storage"),
            ("role", "ENUM", "RoleEnum ('student' | 'admin')"),
            ("organizationId", "UUID", "Foreign Key -> organizations(id)"),
            ("name / contact", "VARCHAR", "User profile details"),
            ("idProofUrl", "VARCHAR(255)", "Secure document link uploaded to Cloudinary"),
            ("volunteer", "BOOLEAN", "Volunteer approval flag (Default False)")
        ]),
        ("3. journal_entries", "Maintains students' personal daily journals and mood metrics.", [
            ("id", "UUID", "Primary Key"),
            ("studentId", "UUID", "Foreign Key -> users(id)"),
            ("content", "TEXT", "Raw journal text content"),
            ("mood_disturbance", "DOUBLE PRECISION", "AI-graded score (0.0 to 10.0)"),
            ("sleep_disruption", "DOUBLE PRECISION", "AI-graded score (0.0 to 10.0)"),
            ("appetite_issues", "DOUBLE PRECISION", "AI-graded score (0.0 to 10.0)"),
            ("academic_disengagement", "DOUBLE PRECISION", "AI-graded score (0.0 to 10.0)"),
            ("social_withdrawal", "DOUBLE PRECISION", "AI-graded score (0.0 to 10.0)"),
            ("date", "TIMESTAMP", "Calendar date of entry (local time alignment)")
        ]),
        ("4. student_moods", "Inference tracking for face/voice checks.", [
            ("id", "UUID", "Primary Key"),
            ("studentId", "UUID", "Foreign Key -> users(id)"),
            ("mood", "ENUM", "MoodEnum ('happy'|'sad'|'angry'|'surprised'|'disgusted'|'fearful'|'neutral')"),
            ("moodScore", "DOUBLE PRECISION", "ML inference confidence rating (0.0 to 1.0)"),
            ("organizationId", "UUID", "Foreign Key -> organizations(id)")
        ]),
        ("5. phq / gad / pss", "Standardised surveys scores.", [
            ("id", "UUID", "Primary Key"),
            ("studentId", "UUID", "Foreign Key -> users(id)"),
            ("score", "INTEGER", "Test points (PHQ-9: 0-27 | GAD-7: 0-21 | PSS-10: 0-40)"),
            ("organizationId", "UUID", "Foreign Key -> organizations(id)"),
            ("takenOn", "TIMESTAMP", "Date of evaluation")
        ]),
        ("6. session_bookings", "Handles counselor session bookings.", [
            ("id", "UUID", "Primary Key"),
            ("studentId", "UUID", "Foreign Key -> users(id)"),
            ("sessionType", "ENUM", "SessionTypeEnum ('individual-counseling' | 'crisis-support' | 'other')"),
            ("reason", "TEXT", "Statement of counseling objective"),
            ("mode", "ENUM", "SessionModeEnum ('virtual' | 'physical')"),
            ("urgency", "ENUM", "UrgencyEnum ('routine' | 'priority' | 'urgent')"),
            ("preferredDate / Time", "DATE / TIME", "Appointment slot preference"),
            ("status", "VARCHAR", "Status of booking ('pending'|'approved'|'completed'|'cancelled'|'rejected')")
        ]),
        ("7. chat_requests", "Pending queue for live peer sessions.", [
            ("id", "UUID", "Primary Key"),
            ("studentId", "UUID", "Foreign Key -> users(id)"),
            ("status", "ENUM", "ChatStatusEnum ('pending' | 'accepted' | 'cancelled')"),
            ("organizationId", "UUID", "Institutional scoping code")
        ]),
        ("8. chat_sessions", "Anonymous peer-support session tunnels.", [
            ("id", "UUID", "Primary Key"),
            ("studentId", "UUID", "Foreign Key -> users(id) (Student)"),
            ("volunteerId", "UUID", "Foreign Key -> users(id) (Volunteer)"),
            ("startedAt / endedAt", "TIMESTAMP", "Session lifetime logging")
        ]),
        ("9. chat_messages", "Temporary messages in anonymous sessions.", [
            ("id", "UUID", "Primary Key"),
            ("chatSessionId", "UUID", "Foreign Key -> chat_sessions(id)"),
            ("senderId", "UUID", "Foreign Key -> users(id)"),
            ("message", "TEXT", "Encrypted or raw text message")
        ])
    ]

    for table_name, desc, fields in tables_spec:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(table_name)
        r_t.bold = True
        r_t.font.size = Pt(12)
        r_t.font.color.rgb = RGBColor(31, 78, 121)
        p_t.add_run(f" — {desc}")
        p_t.paragraph_format.space_before = Pt(8)
        p_t.paragraph_format.space_after = Pt(2)
        
        # Add table
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Light Shading Accent 1'
        hdr_cells = t.rows[0].cells
        hdr_cells[0].text = 'Field Name'
        hdr_cells[1].text = 'Data Type'
        hdr_cells[2].text = 'Constraints & Details'
        
        for c in hdr_cells:
            set_cell_background(c, '1F4E79')
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for fname, ftype, fdetail in fields:
            row_cells = t.add_row().cells
            row_cells[0].text = fname
            row_cells[1].text = ftype
            row_cells[2].text = fdetail
            
            # Format text font sizing inside tables
            for c in row_cells:
                for p in c.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(9.5)

    # Save document
    doc.save("Campus_Care_Documentation.docx")
    print("SUCCESS: Campus_Care_Documentation.docx generated successfully.")

if __name__ == '__main__':
    main()
